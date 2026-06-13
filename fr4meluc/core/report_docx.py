"""Generador de informes DOCX profesionales para Fr4meLuc Enterprise.

Construye un informe de auditoría (portada, resumen ejecutivo, hallazgos,
recomendaciones y anexos) en formato Word usando python-docx.

Decisiones de diseño:
  * El sombreado de celdas se evita en filas de datos: las severidades se
    representan con texto en negrita + prefijo (emoji + etiqueta). El único
    relleno XML se aplica a la fila de cabecera de la tabla.
  * Degradación elegante: si python-docx no está instalado, imprime un mensaje
    y devuelve None sin lanzar excepción.
"""

import os

from .storage import get_scan, get_scan_findings


# ─────────────────────────────────────────────────────────────
#  Constantes de severidad
# ─────────────────────────────────────────────────────────────

_SEV_ORDER = {"critical": 0, "high": 1, "medium": 2, "low": 3, "info": 4}

_SEV_CVSS = {
    "critical": 9.5,
    "high": 8.0,
    "medium": 5.5,
    "low": 2.0,
    "info": 0.0,
}

_SEV_LABEL = {
    "critical": "CRÍTICO",
    "high": "ALTO",
    "medium": "MEDIO",
    "low": "BAJO",
    "info": "INFO",
}

# Prefijo con marcador para la celda de severidad (sin sombreado de celda).
_SEV_PREFIX = {
    "critical": "🔴 CRÍTICO",
    "high": "🟠 ALTO",
    "medium": "🟡 MEDIO",
    "low": "🔵 BAJO",
    "info": "⚪ INFO",
}

# Color RGB (hex) por severidad, usado para texto de severidad coloreado.
_SEV_RGB = {
    "critical": (0x7C, 0x3A, 0xED),
    "high": (0xDC, 0x26, 0x26),
    "medium": (0xD9, 0x77, 0x06),
    "low": (0x25, 0x63, 0xEB),
    "info": (0x6B, 0x72, 0x80),
}

_RECOMMENDATION = {
    "critical": "Remediación inmediata requerida (< 24h). Escalar a CISO.",
    "high": "Remediación urgente (< 72h). Revisar controles de acceso.",
    "medium": "Planificar remediación en el siguiente ciclo (< 30 días).",
    "low": "Incluir en backlog de seguridad (< 90 días).",
    "info": "Revisar como mejora de visibilidad. Sin urgencia operativa.",
}


# ─────────────────────────────────────────────────────────────
#  Helpers de datos (independientes de python-docx)
# ─────────────────────────────────────────────────────────────

def _load_settings() -> dict:
    """Carga `fr4meluc_settings.json` del cwd. Devuelve {} si falta o es inválido."""
    import json
    try:
        with open("fr4meluc_settings.json", "r", encoding="utf-8") as f:
            return json.load(f)
    except (FileNotFoundError, json.JSONDecodeError):
        return {}


def _normalize_severity(value: str) -> str:
    sev = (value or "").strip().lower()
    return sev if sev in _SEV_ORDER else "info"


def _sort_findings(findings: list) -> list:
    return sorted(
        findings,
        key=lambda f: _SEV_ORDER.get(_normalize_severity(f.get("severity")), 4),
    )


def _severity_counts(findings: list) -> dict:
    counts = {sev: 0 for sev in _SEV_ORDER}
    for f in findings:
        counts[_normalize_severity(f.get("severity"))] += 1
    return counts


def _overall_risk(findings: list) -> str:
    if not findings:
        return "info"
    return min(
        (_normalize_severity(f.get("severity")) for f in findings),
        key=lambda s: _SEV_ORDER[s],
    )


def _present_severities(findings: list) -> list:
    present = {_normalize_severity(f.get("severity")) for f in findings}
    return sorted(present, key=lambda s: _SEV_ORDER[s])


def _truncate(text: str, limit: int) -> str:
    text = text or ""
    if len(text) <= limit:
        return text
    return text[:limit].rstrip() + "…"


# ─────────────────────────────────────────────────────────────
#  Helper de sombreado (solo para la fila de cabecera)
# ─────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    """Aplica relleno de color a una celda. Uso reservado a la cabecera."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


# ─────────────────────────────────────────────────────────────
#  Construcción del documento
# ─────────────────────────────────────────────────────────────

def _build_document(scan: dict, findings: list, settings: dict,
                    target: str, domain: str | None):
    """Construye y devuelve el objeto Document de python-docx."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    GRAY = RGBColor(0x55, 0x55, 0x55)
    DARK_BLUE = RGBColor(0x1A, 0x1A, 0x2E)
    RED = RGBColor(0xC0, 0x00, 0x00)

    doc = Document()

    # ── Configuración de página y fuente por defecto ──
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    def _centered(text: str, size: int, bold: bool = False,
                  color: "RGBColor | None" = None):
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        run.font.size = Pt(size)
        run.bold = bold
        if color is not None:
            run.font.color.rgb = color
        return para

    company = settings.get("company_name") or "Fr4meLuc Enterprise"
    auditor = settings.get("auditor_name") or "—"
    auditor_email = settings.get("auditor_email") or "—"
    scan_date = scan.get("started_at") or "—"

    # ── PORTADA ──
    _centered(company, 14, color=GRAY)
    _centered("INFORME DE AUDITORÍA DE SEGURIDAD", 22, bold=True, color=DARK_BLUE)
    _centered(f"Target: {target}", 12)
    _centered(f"Auditor: {auditor}  |  {auditor_email}", 10, color=GRAY)
    _centered(f"Fecha: {scan_date}", 10, color=GRAY)
    _centered("CONFIDENCIAL — NO DISTRIBUIR", 11, bold=True, color=RED)
    doc.add_page_break()

    # ── RESUMEN EJECUTIVO ──
    doc.add_heading("Resumen Ejecutivo", level=1)

    risk = _overall_risk(findings)
    risk_para = doc.add_paragraph()
    risk_run = risk_para.add_run(f"RIESGO GLOBAL: {_SEV_LABEL[risk]}")
    risk_run.bold = True
    risk_run.font.color.rgb = RGBColor(*_SEV_RGB[risk])

    counts = _severity_counts(findings)
    counts_line = "  |  ".join(
        f"{_SEV_LABEL[sev]}: {counts[sev]}" for sev in _SEV_ORDER
    )
    doc.add_paragraph(counts_line)

    disclaimer_text = (
        f"Este informe representa el estado de seguridad del sistema evaluado en el "
        f"momento del análisis ({scan_date}). Los resultados son válidos únicamente "
        f"para el alcance definido: {target}. No constituye una garantía de seguridad "
        f"total del sistema. Fr4meLuc es una herramienta de apoyo para auditorías "
        f"autorizadas."
    )
    disc_para = doc.add_paragraph()
    disc_run = disc_para.add_run(disclaimer_text)
    disc_run.italic = True
    disc_run.font.size = Pt(9)

    # ── HALLAZGOS ──
    doc.add_heading("Hallazgos Detallados", level=1)
    sorted_findings = _sort_findings(findings)

    if sorted_findings:
        headers = ["#", "Título", "Severidad", "CVSS", "Categoría", "Detalle"]
        widths = [Cm(1), Cm(6), Cm(2), Cm(1.5), Cm(2.5), Cm(5)]
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"

        # Cabecera: texto blanco, negrita, fondo azul oscuro (único uso de XML fill).
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            cell = header_cells[i]
            cell.width = widths[i]
            _set_cell_bg(cell, "1A1A2E")
            run = cell.paragraphs[0].add_run(header)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for idx, f in enumerate(sorted_findings, start=1):
            sev = _normalize_severity(f.get("severity"))
            row = table.add_row().cells
            for col, width in enumerate(widths):
                row[col].width = width

            row[0].text = str(idx)
            row[1].text = f.get("title") or "—"

            # Severidad: negrita con prefijo de marcador (sin sombreado de celda).
            sev_run = row[2].paragraphs[0].add_run(_SEV_PREFIX[sev])
            sev_run.bold = True

            row[3].text = f"{_SEV_CVSS[sev]:.1f}"
            row[4].text = f.get("category") or "—"
            row[5].text = _truncate(f.get("detail") or "", 200)
    else:
        doc.add_paragraph("No se registraron hallazgos para este escaneo.")

    # ── RECOMENDACIONES ──
    doc.add_heading("Recomendaciones", level=1)
    present = _present_severities(findings)
    if present:
        for sev in present:
            doc.add_heading(_SEV_LABEL[sev], level=2)
            doc.add_paragraph(_RECOMMENDATION[sev])
    else:
        doc.add_paragraph("Sin hallazgos que requieran recomendaciones.")

    # ── ANEXOS ──
    doc.add_heading("Anexos — Evidencias", level=1)
    workspace_dir = scan.get("workspace_dir") or ""
    txt_files: list[str] = []
    if workspace_dir and os.path.isdir(workspace_dir):
        for root, _dirs, files in os.walk(workspace_dir):
            for name in sorted(files):
                if name.lower().endswith(".txt"):
                    txt_files.append(os.path.join(root, name))
    txt_files.sort()

    if txt_files:
        for path in txt_files:
            rel = os.path.relpath(path, workspace_dir)
            doc.add_heading(rel, level=2)
            try:
                with open(path, "r", encoding="utf-8", errors="replace") as fh:
                    content = fh.read()
            except OSError as exc:
                content = f"[No se pudo leer el archivo: {exc}]"
            body = _truncate(content, 2000)
            para = doc.add_paragraph()
            run = para.add_run(body)
            run.font.name = "Courier New"
            run.font.size = Pt(8)
    else:
        doc.add_paragraph("No se encontraron evidencias adicionales.")

    return doc


# ─────────────────────────────────────────────────────────────
#  API pública
# ─────────────────────────────────────────────────────────────

def generate_docx_report(scan_id: int, workspace_dir: str, target: str,
                         domain: str | None = None) -> str | None:
    """Genera un informe DOCX profesional para `scan_id`.

    Devuelve la ruta al DOCX generado, o None si python-docx no está disponible
    o si ocurre un error durante la generación.
    """
    try:
        import docx  # noqa: F401  (probe de disponibilidad)
    except ImportError:
        print("[ERROR] python-docx no está instalado. Instálalo con: pip install python-docx")
        return None

    try:
        scan = get_scan(scan_id)
        if scan is None:
            print(f"[ERROR] Scan ID {scan_id} no encontrado en la DB.")
            return None

        findings = get_scan_findings(scan_id)
        settings = _load_settings()
        document = _build_document(scan, findings, settings, target, domain)

        output_path = os.path.join(
            workspace_dir, f"Reporte_Pentest_{target.replace('.', '_')}.docx"
        )
        document.save(output_path)
        return output_path
    except Exception as exc:
        print(f"[ERROR] Fallo al generar el DOCX: {exc}")
        return None
